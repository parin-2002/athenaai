from langchain_core.documents import Document
from PyPDF2 import PdfReader
import docx

class DocumentExtractor:
    """
    A class to extract text from PDF and DOCX files and create Document objects.

    Attributes:
        None

    Methods:
        get_text(docs): Extracts text from PDF and DOCX files.
    """

    def __init__(self):
        pass

    def get_text(self, docs):
        """
        Extract text from PDF and DOCX files.

        Args:
            docs (list): A list of file objects (PDF or DOCX).

        Returns:
            list: List of Document objects created from extracted text.
        """
        # Initialize an empty list to store the extracted text
        text = []
        # Initialize an empty list to store the file names
        file_names=[]

        # Iterate through each document in the input list
        for doc in docs:
            # print(doc.name)

            # Extract the file name from the UploadedFile object
            file_name = doc.name if hasattr(doc, 'name') else str(doc)  

            # Check the file extension of the document
            if file_name.endswith('.pdf'):
                pdf_text=""
                # If the file is a PDF
                pdf_reader = PdfReader(doc)
                # Iterate through each page in the PDF
                for page in pdf_reader.pages:
                    # Extract text from the page and append to the pdf_text variable
                    pdf_text += page.extract_text()
                pdf_document=Document(page_content=pdf_text)
                pdf_document.metadata={'Name':doc.name}
                text.append(pdf_document)
                file_names.append(doc.name)


            elif file_name.endswith('.docx'):
                docx_text=""
                # If the file is a DOCX
                docx_file = docx.Document(doc)
                # Iterate through each paragraph in the DOCX
                for para in docx_file.paragraphs:
                    # Extract text from the paragraph and append to the docx_text variable
                    docx_text += para.text
                docx_document=Document(page_content=docx_text)
                docx_document.metadata={'Name':doc.name}
                text.append(docx_document)
                file_names.append(doc.name)
            

        # Return the list of Document objects created from extracted text      
        return text, file_names



### Example of how you can use the DocumentExtractor class

# # Import the DocumentExtractor class
# from document_extractor import DocumentExtractor

# # Create an instance of DocumentExtractor
# doc_extractor = DocumentExtractor()

# # Example PDF and DOCX files
# pdf_file = 'example.pdf'
# docx_file = 'example.docx'

# # Open PDF and DOCX files in binary read mode
# with open(pdf_file, 'rb') as pdf, open(docx_file, 'rb') as docx:
#     # Provide a list of documents to extract text from
#     documents = [pdf, docx]
#     # Call the get_text method to extract text from the documents
#     extracted_documents = doc_extractor.get_text(documents)

# # Print the extracted text from each document
# for document in extracted_documents:
#     print("Document Name:", document.metadata['Name'])
#     print("Extracted Text:", document.page_content)
#     print()
